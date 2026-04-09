"""
Export Service — Generate branded PDF/DOCX audit reports.

Uses ReportLab for PDF and python-docx for DOCX generation.
Includes: cover page, audit summary, risk table, suggestions, and chat history.
"""

import os
import io
from datetime import datetime
from typing import Dict, Any, List, Optional

from reportlab.lib.pagesizes import A4
from reportlab.lib.units import inch, cm
from reportlab.lib.colors import HexColor
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.platypus import (
    SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle,
    PageBreak, HRFlowable
)
from reportlab.lib.enums import TA_CENTER, TA_LEFT

try:
    from docx import Document as DocxDocument
    from docx.shared import Inches, Pt, Cm, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_TABLE_ALIGNMENT
    HAS_DOCX = True
except ImportError:
    HAS_DOCX = False


# Brand colors
PRIMARY_BLUE = HexColor("#2563eb")
DARK_SLATE = HexColor("#0f172a")
SLATE_600 = HexColor("#475569")
RED_500 = HexColor("#ef4444")
AMBER_500 = HexColor("#f59e0b")
GREEN_500 = HexColor("#10b981")
WHITE = HexColor("#ffffff")
LIGHT_BG = HexColor("#f8fafc")


class ExportService:

    def generate_pdf(self, audit_data: dict, chat_history: List[dict] = None) -> io.BytesIO:
        """Generates a branded PDF audit report."""
        buffer = io.BytesIO()
        doc = SimpleDocTemplate(
            buffer, pagesize=A4,
            topMargin=1.5 * cm, bottomMargin=1.5 * cm,
            leftMargin=2 * cm, rightMargin=2 * cm
        )

        styles = getSampleStyleSheet()

        # Custom styles
        title_style = ParagraphStyle(
            'CustomTitle', parent=styles['Title'],
            fontSize=24, spaceAfter=6, textColor=DARK_SLATE, fontName='Helvetica-Bold'
        )
        subtitle_style = ParagraphStyle(
            'CustomSubtitle', parent=styles['Normal'],
            fontSize=11, textColor=SLATE_600, spaceAfter=20
        )
        heading_style = ParagraphStyle(
            'CustomHeading', parent=styles['Heading2'],
            fontSize=14, textColor=PRIMARY_BLUE, spaceBefore=20, spaceAfter=10,
            fontName='Helvetica-Bold'
        )
        body_style = ParagraphStyle(
            'CustomBody', parent=styles['Normal'],
            fontSize=10, textColor=DARK_SLATE, spaceAfter=8, leading=14
        )
        small_style = ParagraphStyle(
            'SmallText', parent=styles['Normal'],
            fontSize=8, textColor=SLATE_600
        )

        elements = []

        # ── Cover Section ──
        elements.append(Spacer(1, 2 * cm))
        elements.append(Paragraph("AUDITOR AI", ParagraphStyle(
            'Brand', parent=styles['Normal'],
            fontSize=12, textColor=PRIMARY_BLUE, fontName='Helvetica-Bold',
            alignment=TA_CENTER, spaceAfter=6
        )))
        elements.append(Paragraph("Audit Report", ParagraphStyle(
            'ReportTitle', parent=styles['Title'],
            fontSize=28, textColor=DARK_SLATE, fontName='Helvetica-Bold',
            alignment=TA_CENTER, spaceAfter=12
        )))
        elements.append(HRFlowable(width="60%", color=PRIMARY_BLUE, thickness=2, spaceAfter=16))

        file_name = audit_data.get("file_name", "Unknown Document")
        created_at = audit_data.get("created_at", datetime.utcnow())
        if isinstance(created_at, str):
            try:
                created_at = datetime.fromisoformat(created_at)
            except:
                created_at = datetime.utcnow()
        date_str = created_at.strftime("%B %d, %Y at %I:%M %p") if created_at else "N/A"

        elements.append(Paragraph(f"<b>Document:</b> {file_name}", ParagraphStyle(
            'MetaInfo', parent=body_style, alignment=TA_CENTER, spaceAfter=4
        )))
        elements.append(Paragraph(f"<b>Generated:</b> {date_str}", ParagraphStyle(
            'MetaInfo2', parent=body_style, alignment=TA_CENTER, spaceAfter=30
        )))

        # ── Risk Score ──
        risk_score = audit_data.get("risk_score", 0)
        score_color = GREEN_500 if risk_score > 80 else AMBER_500 if risk_score >= 50 else RED_500
        score_label = "Safe" if risk_score > 80 else "Moderate Risk" if risk_score >= 50 else "Critical Risk"

        elements.append(Paragraph("Risk Assessment", heading_style))
        score_data = [
            ["Risk Score", "Assessment"],
            [str(risk_score) + " / 100", score_label]
        ]
        score_table = Table(score_data, colWidths=[3 * inch, 3 * inch])
        score_table.setStyle(TableStyle([
            ('BACKGROUND', (0, 0), (-1, 0), PRIMARY_BLUE),
            ('TEXTCOLOR', (0, 0), (-1, 0), WHITE),
            ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
            ('FONTSIZE', (0, 0), (-1, -1), 10),
            ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
            ('VALIGN', (0, 0), (-1, -1), 'MIDDLE'),
            ('GRID', (0, 0), (-1, -1), 0.5, SLATE_600),
            ('ROWBACKGROUNDS', (0, 1), (-1, -1), [LIGHT_BG, WHITE]),
            ('TOPPADDING', (0, 0), (-1, -1), 8),
            ('BOTTOMPADDING', (0, 0), (-1, -1), 8),
        ]))
        elements.append(score_table)

        # ── Summary ──
        elements.append(Paragraph("Executive Summary", heading_style))
        summary = audit_data.get("summary", "No summary available.")
        elements.append(Paragraph(summary, body_style))

        # ── Risks ──
        risks = audit_data.get("risks", [])
        if risks:
            elements.append(Paragraph(f"Identified Risks ({len(risks)})", heading_style))
            risk_table_data = [["#", "Risk Title", "Severity", "Description"]]
            for i, r in enumerate(risks, 1):
                risk_table_data.append([
                    str(i),
                    r.get("title", "Unknown"),
                    r.get("severity", "N/A"),
                    r.get("description", "")[:100] + ("..." if len(r.get("description", "")) > 100 else "")
                ])
            risk_table = Table(risk_table_data, colWidths=[0.5 * inch, 1.5 * inch, 1 * inch, 3 * inch])
            risk_table.setStyle(TableStyle([
                ('BACKGROUND', (0, 0), (-1, 0), RED_500),
                ('TEXTCOLOR', (0, 0), (-1, 0), WHITE),
                ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
                ('FONTSIZE', (0, 0), (-1, -1), 9),
                ('ALIGN', (0, 0), (0, -1), 'CENTER'),
                ('VALIGN', (0, 0), (-1, -1), 'TOP'),
                ('GRID', (0, 0), (-1, -1), 0.5, SLATE_600),
                ('ROWBACKGROUNDS', (0, 1), (-1, -1), [LIGHT_BG, WHITE]),
                ('TOPPADDING', (0, 0), (-1, -1), 6),
                ('BOTTOMPADDING', (0, 0), (-1, -1), 6),
            ]))
            elements.append(risk_table)

        # ── Chat History (if provided) ──
        if chat_history:
            elements.append(PageBreak())
            elements.append(Paragraph("Chat History", heading_style))
            for msg in chat_history:
                role = "You" if msg.get("role") == "user" else "Auditor AI"
                content = msg.get("content", "")
                elements.append(Paragraph(f"<b>{role}:</b>", ParagraphStyle(
                    'ChatRole', parent=body_style, textColor=PRIMARY_BLUE if role == "Auditor AI" else DARK_SLATE
                )))
                elements.append(Paragraph(content, body_style))
                elements.append(Spacer(1, 6))

        # ── Footer ──
        elements.append(Spacer(1, 1 * cm))
        elements.append(HRFlowable(width="100%", color=SLATE_600, thickness=0.5))
        elements.append(Paragraph(
            "Generated by Auditor AI — This report is AI-generated and should not replace professional legal advice.",
            small_style
        ))

        doc.build(elements)
        buffer.seek(0)
        return buffer

    def generate_docx(self, audit_data: dict, chat_history: List[dict] = None) -> io.BytesIO:
        """Generates a branded DOCX audit report."""
        if not HAS_DOCX:
            raise ImportError("python-docx is required for DOCX export")

        doc = DocxDocument()

        # Style setup
        style = doc.styles['Normal']
        style.font.name = 'Calibri'
        style.font.size = Pt(10)

        # Cover
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run("AUDITOR AI")
        run.font.size = Pt(12)
        run.font.color.rgb = RGBColor(0x25, 0x63, 0xeb)
        run.bold = True

        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run("Audit Report")
        run.font.size = Pt(24)
        run.bold = True

        doc.add_paragraph()  # spacer

        file_name = audit_data.get("file_name", "Unknown")
        doc.add_paragraph(f"Document: {file_name}")

        created_at = audit_data.get("created_at", datetime.utcnow())
        if isinstance(created_at, str):
            try:
                created_at = datetime.fromisoformat(created_at)
            except:
                created_at = datetime.utcnow()
        doc.add_paragraph(f"Generated: {created_at.strftime('%B %d, %Y') if created_at else 'N/A'}")

        # Risk Score
        risk_score = audit_data.get("risk_score", 0)
        score_label = "Safe" if risk_score > 80 else "Moderate Risk" if risk_score >= 50 else "Critical Risk"
        doc.add_heading("Risk Assessment", level=2)
        doc.add_paragraph(f"Score: {risk_score}/100 — {score_label}")

        # Summary
        doc.add_heading("Executive Summary", level=2)
        doc.add_paragraph(audit_data.get("summary", "No summary available."))

        # Risks
        risks = audit_data.get("risks", [])
        if risks:
            doc.add_heading(f"Identified Risks ({len(risks)})", level=2)
            table = doc.add_table(rows=1, cols=4)
            table.style = 'Light Grid Accent 1'
            header = table.rows[0].cells
            header[0].text = "#"
            header[1].text = "Risk Title"
            header[2].text = "Severity"
            header[3].text = "Description"
            for i, r in enumerate(risks, 1):
                row = table.add_row().cells
                row[0].text = str(i)
                row[1].text = r.get("title", "Unknown")
                row[2].text = r.get("severity", "N/A")
                row[3].text = r.get("description", "")

        # Chat History
        if chat_history:
            doc.add_page_break()
            doc.add_heading("Chat History", level=2)
            for msg in chat_history:
                role = "You" if msg.get("role") == "user" else "Auditor AI"
                p = doc.add_paragraph()
                run = p.add_run(f"{role}: ")
                run.bold = True
                p.add_run(msg.get("content", ""))

        # Footer
        doc.add_paragraph()
        p = doc.add_paragraph("Generated by Auditor AI — This report is AI-generated and should not replace professional legal advice.")
        p.runs[0].font.size = Pt(8)
        p.runs[0].font.color.rgb = RGBColor(0x64, 0x74, 0x8B)

        buffer = io.BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        return buffer
